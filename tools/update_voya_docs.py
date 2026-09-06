from pathlib import Path
from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r'C:\wamp64\www\voya-vendor-platform')
SOURCE = ROOT / 'Voya_Activity_SaaS_User_Guide_Admin_Flow_Updated.docx'
OUTPUT = ROOT / 'Voya_Activity_SaaS_User_Guide_Final_With_ERD.docx'
QA = ROOT / 'docx_qa'
QA.mkdir(exist_ok=True)
ERD = QA / 'voya_database_erd.png'

font_path = r'C:\Windows\Fonts\segoeui.ttf'
bold_path = r'C:\Windows\Fonts\segoeuib.ttf'
font = ImageFont.truetype(font_path, 22)
small = ImageFont.truetype(font_path, 17)
title_font = ImageFont.truetype(bold_path, 26)

models = {
    'Tenant': ['id PK', 'name', 'slug UQ', 'kind'],
    'User': ['id PK', 'tenantId FK', 'email UQ', 'role', 'active'],
    'VendorProfile': ['id PK', 'tenantId FK UQ', 'business fields', 'payout fields', 'documentStatus JSON'],
    'Activity': ['id PK', 'tenantId FK', 'product fields', 'location fields', 'status'],
    'ActivityMedia': ['id PK', 'activityId FK', 'kind', 'url', 'description', 'rank'],
    'RatePlan': ['id PK', 'activityId FK', 'pricing fields', 'traveller fields', 'policy fields'],
    'TravellerRule': ['id PK', 'ratePlanId FK', 'type', 'age/count', 'price'],
    'CancellationRule': ['id PK', 'ratePlanId FK', 'days', 'charge'],
    'AvailabilitySlot': ['id PK', 'ratePlanId FK', 'slotDate', 'startTime', 'capacity', 'available'],
    'Promotion': ['id PK', 'activityId FK', 'discount', 'window', 'active'],
    'Booking': ['id PK', 'tenantId FK', 'activityId FK', 'ratePlanId FK', 'bookingCode UQ', 'status'],
    'Payout': ['id PK', 'tenantId FK', 'amount', 'status', 'dueDate'],
    'AuditLog': ['id PK', 'tenantId FK', 'actorId', 'action', 'entityType', 'payload JSON'],
}
positions = {
    'Tenant': (50, 140), 'User': (430, 140), 'VendorProfile': (430, 330), 'Activity': (800, 140),
    'ActivityMedia': (1180, 40), 'RatePlan': (1180, 300), 'TravellerRule': (1550, 220),
    'CancellationRule': (1550, 470), 'AvailabilitySlot': (1550, 720), 'Promotion': (800, 500),
    'Booking': (430, 610), 'Payout': (50, 700), 'AuditLog': (50, 930),
}
box_w = 330
box_h = 150
img = Image.new('RGB', (1960, 1160), '#F8FAFC')
draw = ImageDraw.Draw(img)
draw.text((50, 25), 'Voya Activity SaaS PostgreSQL ERD', fill='#14213D', font=title_font)
draw.text((50, 62), 'Generated from the current Prisma schema and verified against database row counts', fill='#52627D', font=small)

for name, (x, y) in positions.items():
    draw.rounded_rectangle((x, y, x + box_w, y + box_h), radius=12, fill='#FFFFFF', outline='#CBD5E1', width=2)
    draw.rectangle((x, y, x + box_w, y + 34), fill='#14213D')
    draw.text((x + 12, y + 6), name, fill='#FFFFFF', font=font)
    for i, field in enumerate(models[name]):
        draw.text((x + 14, y + 43 + i * 19), field, fill='#334155', font=small)

relations = [
    ('Tenant', 'User'), ('Tenant', 'VendorProfile'), ('Tenant', 'Activity'), ('Tenant', 'Booking'), ('Tenant', 'Payout'), ('Tenant', 'AuditLog'),
    ('Activity', 'ActivityMedia'), ('Activity', 'RatePlan'), ('Activity', 'Promotion'), ('Activity', 'Booking'),
    ('RatePlan', 'TravellerRule'), ('RatePlan', 'CancellationRule'), ('RatePlan', 'AvailabilitySlot'), ('RatePlan', 'Booking'),
]
for a, b in relations:
    ax, ay = positions[a]; bx, by = positions[b]
    start = (ax + box_w, ay + box_h // 2) if ax < bx else (ax, ay + box_h // 2)
    end = (bx, by + box_h // 2) if ax < bx else (bx + box_w, by + box_h // 2)
    draw.line((start, end), fill='#94A3B8', width=2)
    draw.ellipse((end[0] - 4, end[1] - 4, end[0] + 4, end[1] + 4), fill='#E8A33D')
img.save(ERD, optimize=True)

doc = Document(SOURCE)
doc.add_page_break()
doc.add_heading('Implementation Update September 2026', level=1)
doc.add_paragraph('This appendix records the completed implementation and the verification performed against the running Chrome application at http://localhost:3007/ and the PostgreSQL database activity_saas. It supplements the existing vendor and administrator flow documentation.')

doc.add_heading('Completed Vendor Features', level=2)
for text in [
    'Listings now support the prototype-style My Listings and Add / Edit Activity flow with seven linked sections: Basic Info, Category and Location, Media, Logistics and Inclusions, Rate Plan and Travellers, Policies and Cancellation, and Availability and Promo.',
    'Activity fields are saved through the create and update APIs. Rate plan fields, traveller rules, cancellation rules, capacity, availability, and promotions are persisted through their respective APIs.',
    'Cover images and multi-image galleries upload as files into activity-saas-backend/public/uploads. Each file is referenced by an ActivityMedia row containing kind, URL, description, and rank. The editor displays image previews and video controls.',
    'Onboarding Bank and Payout now captures masked payout account, account holder, bank, branch, IFSC, SWIFT or BIC, account type, and payout currency. These values are stored on VendorProfile. Raw account credentials are not collected.',
    'Confirmed booking Voucher buttons generate a real PDF from the booking record, save it under public/uploads, return the API URL, and display the PDF in the application preview modal.',
    'The application main panel owns vertical scrolling while the vendor sidebar remains fixed. This keeps all long screens reachable, including availability inventory and pricing rules.',
]: doc.add_paragraph(text, style='List Bullet')

doc.add_heading('Administrator Flow', level=2)
doc.add_paragraph('An administrator signs in with the administrator account, views all vendor tenants, selects a vendor, reviews uploaded documents, opens image or PDF previews, verifies or rejects individual documents, approves or suspends the vendor, and reviews vendor activities. Activity review supports publishing or rejecting activities after vendor submission. These actions call the admin API and update PostgreSQL records; they are not static screen-only controls.')

doc.add_heading('Chrome Verification Evidence', level=2)
evidence = [
    ('Bank and Payout', 'Chrome displayed the eight payout fields. Save and Continue moved onboarding to Catalogue Setup, and a subsequent API read returned the saved account holder, bank, branch, IFSC, SWIFT, CURRENT account type, and INR currency.'),
    ('Media', 'Chrome displayed Cover Image, Gallery, and Video Link cards. An uploaded image was returned by the API, stored in public/uploads, referenced in ActivityMedia, and rendered as an image preview.'),
    ('Availability', 'Chrome displayed the promotion card, live-deal summary, weekday date headers, slot inventory grid, Bulk Edit, Save Inventory, and Standing Pricing Rules. The slot and promotion screens loaded persisted records.'),
    ('Bookings', 'Chrome displayed confirmed bookings. Clicking Voucher opened voucher-BK-77291.pdf inside the application PDF viewer; the server returned HTTP 200 with content type application/pdf.'),
]
table = doc.add_table(rows=1, cols=2)
table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Screen'
table.rows[0].cells[1].text = 'Verified behavior'
for a, b in evidence:
    cells = table.add_row().cells; cells[0].text = a; cells[1].text = b

doc.add_heading('Database ERD', level=1)
doc.add_paragraph('The following ERD represents the current PostgreSQL structure. Foreign keys are shown by the relationship lines. Tenant is the ownership boundary for vendor data; activities own media, rate plans, promotions, and bookings; rate plans own traveller rules, cancellation rules, and availability slots.')
doc.add_picture(str(ERD), width=Inches(7.0))
doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('Figure 1. Current Voya Activity SaaS database relationship diagram.', style='Caption')

doc.add_heading('Verified Database Snapshot', level=2)
counts = [('Tenant', 12), ('User', 13), ('VendorProfile', 12), ('Activity', 22), ('ActivityMedia', 4), ('RatePlan', 28), ('TravellerRule', 8), ('CancellationRule', 9), ('AvailabilitySlot', 41), ('Promotion', 1), ('Booking', 5), ('Payout', 3), ('AuditLog', 0)]
table = doc.add_table(rows=1, cols=3); table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Table / model'; table.rows[0].cells[1].text = 'Rows'; table.rows[0].cells[2].text = 'Purpose'
purpose = {'Tenant':'Tenant ownership and platform or vendor boundary','User':'Login identities and roles','VendorProfile':'Vendor onboarding, verification and payout metadata','Activity':'Vendor product master','ActivityMedia':'Uploaded or linked media','RatePlan':'Bookable commercial offer','TravellerRule':'Age, count and price rules','CancellationRule':'Cancellation charge rules','AvailabilitySlot':'Date and slot inventory','Promotion':'Time-boxed discounts','Booking':'Reservations and voucher source data','Payout':'Settlement records','AuditLog':'Audit event store'}
for name, n in counts:
    cells = table.add_row().cells; cells[0].text = name; cells[1].text = str(n); cells[2].text = purpose[name]

doc.add_heading('Important Persistence Rules', level=2)
for text in [
    'Vendor records are isolated by tenantId for vendor users. Platform administrators can review across vendor tenants.',
    'Uploaded documents and activity media store file references in PostgreSQL while binary files are served from public/uploads through the API static route.',
    'A vendor edit to a LIVE activity moves the activity back to UNDER_REVIEW so production changes are reviewed before publication.',
    'Availability updates use optimistic version fields and bulk persistence. Promotions store their activity, discount, time window, booking cap, and active state.',
]: doc.add_paragraph(text, style='List Bullet')

doc.save(OUTPUT)
print(OUTPUT)
