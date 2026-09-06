from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(r'C:\wamp64\www\voya-vendor-platform')
ASSETS = ROOT / 'docx_qa' / 'fresh_20260906'
OUT = ROOT / 'Voya_Activity_Marketplace_Client_Ready_Proposal_Rev2.docx'

def shade(cell, fill):
    pr = cell._tc.get_or_add_tcPr(); el = pr.find(qn('w:shd'))
    if el is None: el = OxmlElement('w:shd'); pr.append(el)
    el.set(qn('w:fill'), fill)

def border(cell):
    pr = cell._tc.get_or_add_tcPr(); bs = pr.first_child_found_in('w:tcBorders')
    if bs is None: bs = OxmlElement('w:tcBorders'); pr.append(bs)
    for edge in ('top','left','bottom','right','insideH','insideV'):
        tag = qn('w:' + edge); el = bs.find(tag)
        if el is None: el = OxmlElement('w:' + edge); bs.append(el)
        el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); el.set(qn('w:color'),'D9D9D9')

def cell_text(cell, value, bold=False, color=None, size=9):
    cell.text = ''; p = cell.paragraphs[0]; p.paragraph_format.space_after = Pt(2)
    r = p.add_run(str(value)); r.bold = bold; r.font.size = Pt(size)
    if color: r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER; border(cell)

def make_table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers)); t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    for i,h in enumerate(headers): cell_text(t.rows[0].cells[i], h, True, 'FFFFFF'); shade(t.rows[0].cells[i], '14213D')
    for ri,row in enumerate(rows):
        cells = t.add_row().cells
        for i,value in enumerate(row):
            cell_text(cells[i], value)
            if ri % 2: shade(cells[i], 'F4F6F8')
    doc.add_paragraph().paragraph_format.space_after = Pt(1)

def h(doc, text, level=1):
    p = doc.add_heading(text, level=level); p.paragraph_format.keep_with_next = True; return p

def bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet'); p.paragraph_format.space_after = Pt(3); p.add_run(text)

def numbered(doc, text):
    p = doc.add_paragraph(style='List Number'); p.paragraph_format.space_after = Pt(3); p.add_run(text)

def figure(doc, filename, caption):
    path = ASSETS / filename
    if path.exists():
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_together = True; p.paragraph_format.keep_with_next = True; p.add_run().add_picture(str(path), width=Inches(6.10))
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; p.paragraph_format.keep_together = True; r = p.add_run(caption); r.italic = True; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(85,85,85)

doc = Document(); s = doc.sections[0]
s.top_margin = Inches(.65); s.bottom_margin = Inches(.65); s.left_margin = Inches(.78); s.right_margin = Inches(.78)
doc.styles['Normal'].font.name = 'Aptos'; doc.styles['Normal'].font.size = Pt(10); doc.styles['Normal'].font.color.rgb = RGBColor(35,35,35); doc.styles['Normal'].paragraph_format.space_after = Pt(7); doc.styles['Normal'].paragraph_format.line_spacing = 1.08
for name,size in [('Title',26),('Heading 1',17),('Heading 2',12.5)]:
    st=doc.styles[name]; st.font.name='Aptos'; st.font.size=Pt(size); st.font.bold=True; st.font.color.rgb=RGBColor(0,0,0)

p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('VOYA'); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=RGBColor(232,163,61)
p=doc.add_paragraph(style='Title'); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; p.add_run('Activity Marketplace Vendor Platform')
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Updated Solution Proposal'); r.bold=True; r.font.size=Pt(15)
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('A client-facing operating platform for activity suppliers and marketplace governance'); r.font.size=Pt(11); r.font.color.rgb=RGBColor(80,80,80)
make_table(doc,['Proposal item','Current proposal'],[
    ['Customer / Product Owner','Voya activity marketplace operations'],
    ['Proposed product','Voya Vendor Platform'],
    ['Implemented foundation','React vendor console, NestJS API and PostgreSQL persistence'],
    ['Operating model','Multi-tenant vendor marketplace with administrator review'],
])
p=doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Prepared from the current source code, PostgreSQL schema and running Chrome screens.'); r.italic=True; r.font.size=Pt(9)
doc.add_page_break()

h(doc,'1. Executive Summary')
doc.add_paragraph('Voya is being developed as a multi-tenant activity marketplace platform connecting activity vendors with a central administrative operation. The current implementation provides a working vendor console, administrator console, PostgreSQL-backed data model and NestJS APIs for onboarding, product catalogue, rate plans, availability, promotions, bookings, vouchers and payouts.')
doc.add_paragraph('The intended lifecycle is complete and operational: a new vendor registers, completes onboarding, uploads verification documents, adds payout information, creates activities and rate plans, configures media and inventory, submits listings for review, and receives bookings after approval. Administrators can view all vendor tenants, review documents, approve or suspend vendors, and publish or reject activities.')
for x in ['Reduce manual coordination between vendors and platform operations.','Keep vendor data isolated by tenant while allowing administrators controlled cross-tenant visibility.','Persist business fields, uploaded files, media references, inventory, booking status and workflow state.','Complete the remaining Activity Product Master fields in the Add/Edit experiences.']: bullet(doc,x)

h(doc,'1.1 What the client can demonstrate')
doc.add_paragraph('A client demonstration can follow one continuous story: create or sign in as a vendor, complete the onboarding checklist, add a bookable activity, attach a rate plan and media, open inventory, then switch to the administrator role to review the vendor and publish the activity. The vendor can then manage the resulting booking and open its voucher. This is the core business loop represented by the current application.')
figure(doc,'02_vendor_dashboard.png','Figure 1. Fresh vendor dashboard captured from the running Chrome application on 06 September 2026.')

h(doc,'2. Current Business Need')
doc.add_paragraph('Activity suppliers need a dependable workspace to maintain business identity, compliance documents, activity products, commercial rate plans, media, inventory and booking operations. Platform administrators need a separate control surface to verify vendor readiness and control which activities become live.')
h(doc,'2.1 Vendor-side operational requirements',2)
for x in ['Capture business details, documents and payout details in one onboarding flow.','Maintain product content including logistics, traveller rules, media, cancellation rules and redemption information.','Maintain availability at slot and date level with capacity, remaining inventory, price overrides and closures.','Manage booking confirmation, cancellation and voucher generation.']: bullet(doc,x)
h(doc,'2.2 Platform-side control requirements',2)
for x in ['View all vendor tenants and onboarding status.','Inspect uploaded documents before verification or rejection.','Persist vendor approval, suspension, document review and activity publication actions.','Prevent vendor users from accessing another tenant’s data.']: bullet(doc,x)

h(doc,'3. Proposed Voya Solution')
doc.add_paragraph('The solution is a modular web application with two role-oriented experiences. Vendors manage their own operating data. Administrators manage marketplace readiness and catalogue governance. Both experiences use the same authenticated API and PostgreSQL data model.')
make_table(doc,['Area','Implemented responsibility'],[
 ['Vendor console','Onboarding, listings, rate plans, media, availability, promotions, bookings, vouchers, payouts and performance.'],
 ['Admin console','Dashboard, vendor list, vendor detail, document preview, document review, vendor approval/suspension and activity review.'],
 ['API layer','NestJS modules for authentication, vendor profile, activities, rate plans, availability, bookings, payouts, dashboard and admin operations.'],
 ['Persistence','PostgreSQL via Prisma with tenant relationships, status fields and uploaded-file references.'],
 ['File storage','Documents, activity media and voucher PDFs are stored under backend public/uploads and served through API URLs.'],
])

h(doc,'4. Vendor Onboarding Workflow')
doc.add_paragraph('Onboarding is the entry point for a fresh vendor. Registration creates a vendor tenant and user account. The vendor then completes a five-step readiness journey before the account is ready for platform review.')
for x in ['Register or sign in as a vendor; the account is associated with a vendor tenant and profile.','Complete Business Details: legal business name, operating city and region, GSTIN and category.','Upload GSTIN certificate, PAN card, cancelled cheque or bank proof, and trade licence or activity permit.','Complete Bank and Payout details, including masked account, holder, bank, branch, IFSC, SWIFT, account type and currency.','Create at least one activity and one rate plan, then configure catalogue readiness.','Review readiness and submit the vendor’s activity or onboarding state for administrator action.']: numbered(doc,x)
figure(doc,'01_vendor_onboarding.png','Figure 2. Fresh vendor onboarding screen captured from the running Chrome application on 06 September 2026.')

h(doc,'5. Product Catalogue and Activity Management')
doc.add_paragraph('Activities are the vendor’s bookable products. The Listings workspace is the vendor’s catalogue home: it shows each activity, its lifecycle state, location, category and headline rate, and provides the entry point to Add or Edit Activity. The Activity Editor follows the prototype’s seven-part structure and saves core product content to PostgreSQL. A new activity is first saved as a draft, after which rate plans and media can be attached.')
make_table(doc,['Catalogue section','Current behavior'],[
 ['Basic information','Product name, type, subtype, descriptions, category, rating, highlights, terms and supporting information.'],
 ['Category and location','City, state, country, address, latitude and longitude.'],
 ['Media','Cover image, multi-image gallery, image preview, video URL, description and persisted media URL.'],
 ['Logistics and inclusions','Things to carry and important information; operational logistics are primarily rate-plan specific.'],
 ['Rate plan and travellers','Price, validity, pax, traveller prices, pickup, vehicle, timing, voucher, confirmation and rules.'],
 ['Policies and availability','Terms, structured cancellation rules, inventory slots and promotions.'],
])
figure(doc,'03_vendor_listings.png','Figure 3. Fresh vendor Listings workspace captured from the running Chrome application on 06 September 2026. The Add or Edit Activity entry point is available in the same screen.')

h(doc,'6. Rate Plan and Traveller Rules')
doc.add_paragraph('A rate plan is the commercial and operational version of an activity. One activity can have multiple rate plans with different prices, validity periods, traveller pricing, pickup rules, cancellation charges and confirmation behavior.')
for x in ['Rate-plan code, name, status, currency, unit type and base price.','Validity dates, minimum and maximum pax, valid weekdays and blackout dates.','Duration, meal, pickup, drop-off, vehicle, ticket, entry-fee, voucher, confirmation and auto-redeem settings.','Traveller rules with age range, counts, display name, description and price.','Cancellation rules stored as structured date ranges with percentage or absolute charges.']: bullet(doc,x)
doc.add_paragraph('The current foundation persists the major rate-plan structures. Full completion against the Activity Product Master requires exposing every detailed field in the Add/Edit screens, including Youth, Infant and Group traveller configurations, flexible cancellation slabs, sightseeing and affiliate fields, meeting-point details, vendor codes and other operational metadata.')

h(doc,'7. Availability and Promotion Management')
doc.add_paragraph('Availability is maintained against a selected rate plan. Vendors can work with a date and slot grid, update capacity and available inventory, apply a price override, close a slot and save changes in bulk. The API validates ownership and prevents available inventory from exceeding capacity.')
figure(doc,'04_vendor_availability.png','Figure 4. Fresh availability and pricing screen captured from the running Chrome application on 06 September 2026.')
for x in ['Rate-plan selection determines which inventory is being edited.','Each slot stores date, start time, capacity, available quantity, optional price override and closed state.','Version fields protect updates from overwriting a newer change.','Promotions store name, discount percentage, start/end window, booking cap and active state.']: bullet(doc,x)

h(doc,'8. Booking and Voucher Workflow')
doc.add_paragraph('Bookings are tenant-scoped and displayed to the vendor with filters for all, needs action, confirmed and cancelled states. A pending booking can be confirmed or cancelled. Confirmed and completed bookings can generate a voucher PDF.')
for x in ['A booking is created for an activity and optional rate plan with service date, pax, amount and customer details.','The vendor reviews the booking and confirms or cancels it through the API.','For eligible bookings, the vendor selects Voucher.','The backend generates a PDF, stores it in public/uploads and returns the file URL.','The frontend opens the returned PDF in an in-app preview frame.']: numbered(doc,x)
figure(doc,'05_vendor_bookings.png','Figure 5. Fresh booking list and voucher action captured from the running Chrome application on 06 September 2026.')

h(doc,'9. Administrator Workflow')
doc.add_paragraph('The administrator console is a separate role-protected experience for marketplace governance. It provides cross-tenant visibility that vendor users do not receive.')
h(doc,'9.1 Administrator dashboard',2); figure(doc,'06_admin_dashboard.png','Figure 6. Fresh administrator dashboard captured from the running Chrome application on 06 September 2026.')
for x in ['Summarises vendors, pending vendors, activities, review activities, bookings and pending bookings.','Provides the starting point for operational review.']: bullet(doc,x)
h(doc,'9.2 Vendor review',2); figure(doc,'07_admin_vendors.png','Figure 7. Fresh administrator vendor list captured from the running Chrome application on 06 September 2026.')
figure(doc,'07_admin_vendor_detail.png','Figure 8. Fresh administrator vendor detail and document review captured from the running Chrome application on 06 September 2026.')
for x in ['Lists vendor tenants and verification state.','Opens vendor details with business information, users, listings and rate plans.','Allows vendor approval or suspension.','Shows required document filename and status.','Provides preview for uploaded images and PDFs.','Allows Verify or Reject actions that update persisted document status and review time.']: bullet(doc,x)
h(doc,'9.3 Activity review',2); figure(doc,'08_admin_activity_review.png','Figure 9. Fresh administrator activity review queue captured from the running Chrome application on 06 September 2026.')
for x in ['Lists activities in review with owning vendor and rate-plan count.','Allows publish or reject actions that change activity lifecycle status.']: bullet(doc,x)

h(doc,'10. Technical Architecture')
make_table(doc,['Layer','Implementation'],[
 ['Frontend','React with TypeScript, Vite and role-aware navigation for vendor and administrator consoles.'],
 ['Backend','NestJS modules, controllers, services, DTO validation and JWT authentication.'],
 ['Database','PostgreSQL accessed through Prisma ORM.'],
 ['Authorization','JWT authentication, role guards and tenant ownership checks.'],
 ['Files','Backend public/uploads directory with URL references returned by APIs.'],
 ['Deployment shape','Separate frontend and backend processes with API base URL configuration.'],
])

h(doc,'11. Core Data Model')
doc.add_paragraph('The database models the marketplace around tenants. A tenant owns a vendor profile, activities, bookings and payouts. Activities own media, rate plans and promotions. Rate plans own traveller rules, cancellation rules and availability slots.')
make_table(doc,['Entity','Purpose'],[
 ['Tenant and User','Tenant isolation, account identity, role and active state.'], ['VendorProfile','Business identity, verification, readiness, payout details and document references.'], ['Activity','Product master content, location, status and commercial relationships.'], ['ActivityMedia','Image or video URL, description and rank.'], ['RatePlan','Bookable price and operating rules.'], ['TravellerRule / CancellationRule','Traveller pricing and cancellation windows.'], ['AvailabilitySlot / Promotion','Inventory and time-boxed discount configuration.'], ['Booking / Payout / AuditLog','Booking lifecycle, settlement records and action history structure.'],
])

h(doc,'12. Activity Product Master Alignment')
doc.add_paragraph('The supplied Activity Product Master workbook is treated as the functional requirements baseline for the activity catalogue and rate-plan experience. The current implementation covers the core data model and workflows, but the user-facing forms do not yet expose every detailed workbook field.')
make_table(doc,['Requirement group','Current position','Completion needed'],[
 ['Activity content','Core fields implemented and persisted.','Expose FAQs, affiliates, visibility, amenities, persuasions, labels, redemption, safety, additional information and metaname.'],
 ['Location','Core location fields implemented.','Add structured meeting-point details and city-code handling.'],
 ['Media','Upload, preview, gallery and references implemented.','Add explicit rank/reorder and complete metadata editing.'],
 ['Rate plans','Core price, validity, logistics, traveller and cancellation structures implemented.','Expose all workbook fields and support full edit workflow.'],
 ['Traveller types','Adult, Child and Senior are currently generated.','Add configurable Youth, Infant and Group rules.'],
 ['Dictionary values','Some enum validation exists.','Connect controlled values to dropdowns and validate consistently.'],
])

h(doc,'13. Proposed Completion Plan')
h(doc,'13.1 Workbook-aligned catalogue completion',2)
for x in ['Complete every Activity Product Master field in the Activity Editor.','Complete every rate-plan field and implement Edit, Duplicate and Deactivate actions.','Add configurable traveller categories and flexible cancellation rules.','Add structured meeting point, media ranking and gallery ordering.','Connect Dictionary values to validated dropdown controls.']: bullet(doc,x)
h(doc,'13.2 Operational hardening',2)
for x in ['Add comprehensive audit logging for administrator and vendor actions.','Add stronger file validation, retention and production object-storage option.','Add automated API and browser regression tests for every critical workflow.','Add reporting and operational monitoring as requirements mature.']: bullet(doc,x)

h(doc,'14. Assumptions and Dependencies')
for x in ['The Activity Product Master workbook is the requirements reference, not an import file.','Business owners will confirm definitions for ambiguous fields such as visibility bit, salience, vendor voucher flag and transfer type.','Production file storage, malware scanning, retention and access policies will be confirmed before deployment.','Banking credentials remain masked and raw secrets are not stored in the application database.','Administrator review rules and the exact conditions for vendor verification and activity publication will be confirmed for production.']: bullet(doc,x)

h(doc,'15. Initial Scope Exclusions')
for x in ['Customer-facing marketplace booking and payment collection are outside the current vendor-console scope unless separately commissioned.','Live supplier, hotel, DMC or channel-manager integrations are not part of the current foundation.','Automated settlement reconciliation and full accounting are not yet implemented.','Native mobile applications are not included in this web-platform proposal.']: bullet(doc,x)

h(doc,'16. Conclusion')
doc.add_paragraph('The Voya Vendor Platform has progressed beyond a static prototype: it has a working vendor and administrator operating model backed by NestJS APIs, PostgreSQL persistence and real file references. The foundation supports the core lifecycle from vendor onboarding through catalogue setup, availability, bookings, vouchers and administrative review.')
doc.add_paragraph('The next material step is workbook-aligned completion of the Activity Editor and Rate Plan Editor. That work will make every required product-master field visible, editable, validated and persisted, completing the transition from a strong operational foundation to a full activity product-management platform.')

for section in doc.sections:
    p=section.footer.paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run('Voya Activity Marketplace Proposal'); r.font.size=Pt(8); r.font.color.rgb=RGBColor(100,100,100)
doc.save(OUT); print(OUT)
