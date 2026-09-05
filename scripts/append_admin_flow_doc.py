from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pathlib import Path

root = Path(__file__).resolve().parents[1]
source = root / 'Voya_Activity_SaaS_User_Guide_Complete_Technical_Handover.docx'
output = root / 'Voya_Activity_SaaS_User_Guide_Admin_Flow_Updated.docx'
assets = root / 'scripts' / 'guide_assets'
doc = Document(str(source))

def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr(); shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), fill); tcPr.append(shd)

doc.add_page_break()
doc.add_heading('Admin Console Workflow', level=1)
doc.add_paragraph('This section documents the administrator path now implemented in the running Voya application at http://localhost:3007/. It explains how a platform administrator supervises every vendor tenant, reviews onboarding evidence, and controls catalogue publication. The screenshots below were captured from the live Chrome session after signing in with the seeded admin account.')
doc.add_heading('Admin Login', level=2)
doc.add_paragraph('Use admin@voya.demo with password Demo@123. The API authenticates the account as ADMIN with no tenantId. That role distinction is important: the administrator is platform-scoped, while a VENDOR user is tenant-scoped. After authentication, the application selects the Admin Console shell and does not call vendor-only profile endpoints.')
doc.add_heading('Admin Navigation', level=2)
doc.add_paragraph('The admin menu contains three operational areas. Dashboard provides cross-tenant metrics and workflow entry points. Vendors lists every vendor tenant and opens a vendor detail review. Activity Review exposes draft and under-review activities that are waiting for a publish decision.')

def figure(path, caption):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(assets / path), width=Inches(6.35))
    cp = doc.add_paragraph(caption); cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cp.runs[0].italic = True; cp.runs[0].font.size = Pt(9)

figure('admin_dashboard.png', 'Figure A1. Admin Dashboard captured from the running Chrome application.')
doc.add_paragraph('The dashboard currently reports 4 vendor tenants, 3 vendors pending onboarding review, 6 activities, 1 activity already under review, 5 bookings, and 2 pending bookings. These values are loaded from the admin API rather than hardcoded in the page. The workflow buttons link directly to the Vendors and Activity Review screens.')

doc.add_heading('Vendor List and Onboarding Review', level=2)
figure('admin_vendor_detail.png', 'Figure A2. Vendor list and selected vendor detail captured from the running Chrome application.')
doc.add_paragraph('The Vendors screen loads all tenants whose kind is VENDOR. Search filters the list by business name, tenant name, or account email. Selecting a vendor loads its complete review context: business profile, operating location, account owner, verification status, uploaded document metadata, listings, and rate-plan counts.')
doc.add_paragraph('The administrator can approve a vendor, suspend a vendor, verify or reject each document, and inspect the vendor catalogue. Document decisions update VendorProfile.documentStatus with the decision and reviewedAt timestamp. Vendor verification updates VendorProfile.verificationStatus. These actions are protected by JwtAuthGuard and RolesGuard and accept only ADMIN or SUB_ADMIN roles.')

doc.add_heading('Activity Review and Publishing', level=2)
figure('admin_review.png', 'Figure A3. Activity Review queue captured from the running Chrome application.')
doc.add_paragraph('The review queue returns vendor activities in DRAFT or UNDER_REVIEW status. Each row identifies the activity, owning vendor, city, current status, and number of configured rate plans. Publish changes the activity to LIVE, making it eligible for the customer catalogue. Reject changes it to INACTIVE, which keeps it out of the live catalogue until the vendor corrects and resubmits it.')

doc.add_heading('End to End Logical Flow', level=2)
steps = [
    ('1. Vendor registration', 'A new vendor creates an account. The backend creates a new VENDOR tenant, user, and pending VendorProfile with empty document status.'),
    ('2. Vendor onboarding', 'The vendor completes business details, uploads document metadata, configures payout information, creates activities, and adds rate plans.'),
    ('3. Administrative verification', 'The administrator opens Vendors, checks the profile and documents, then verifies or rejects evidence and approves or suspends the vendor.'),
    ('4. Catalogue review', 'The administrator opens Activity Review and checks the activity content, vendor ownership, status, and rate-plan readiness.'),
    ('5. Publication', 'Publishing changes the activity status to LIVE. The vendor can then manage availability and bookings for the published catalogue item.'),
    ('6. Ongoing operations', 'The admin dashboard remains platform-wide, while vendor pages continue to enforce tenant isolation for listings, availability, bookings, payouts, and performance.'),
]
table = doc.add_table(rows=1, cols=2); table.style = 'Table Grid'
table.rows[0].cells[0].text = 'Stage'; table.rows[0].cells[1].text = 'What happens'
for c in table.rows[0].cells: shade(c, '14213D')
for left, right in steps:
    cells = table.add_row().cells; cells[0].text = left; cells[1].text = right

doc.add_heading('Admin API Mapping', level=2)
doc.add_paragraph('The UI is backed by these protected endpoints: GET /api/admin/dashboard, GET /api/admin/vendors, GET /api/admin/vendors/:tenantId, PATCH /api/admin/vendors/:tenantId/verification, PATCH /api/admin/vendors/:tenantId/documents/:key, GET /api/admin/activities/review, POST /api/admin/activities/:id/publish, and POST /api/admin/activities/:id/reject. All endpoints require a JWT and an ADMIN or SUB_ADMIN role.')
doc.add_heading('Demo Verification Result', level=2)
doc.add_paragraph('Chrome verification completed successfully using admin@voya.demo. The tested path was Admin Dashboard, Vendors, selected vendor detail, and Activity Review. Backend and frontend production builds also completed successfully.')

doc.save(str(output))
print(output)
